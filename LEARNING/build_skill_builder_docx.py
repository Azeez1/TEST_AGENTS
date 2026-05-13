"""Generate Block 2 (the working meta-prompt) as a plain professional .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0x00, 0x00, 0x00)
RED = RGBColor(0x8A, 0x00, 0x00)
GRAY = RGBColor(0x55, 0x55, 0x55)
RULE_GRAY = "888888"

doc = Document()

section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.3

def heading(text, level=1):
    sizes = {1: 22, 2: 16, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes[level])
    r.font.color.rgb = BLACK
    return p

def para(text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color or BLACK
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
    r.font.color.rgb = BLACK
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), RULE_GRAY)
    pBdr.append(left)
    pPr.append(pBdr)
    return p

def rule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════ COVER ══════════════════
heading("Interview-Driven Skill Builder", level=1)
para("Working Meta-Prompt — Operational Version", italic=True, color=GRAY, size=12)
rule()

heading("How to use this document", level=3)
para("This document contains the complete working meta-prompt referenced in the Prompt Library submission titled \"Interview-Driven Skill Builder.\" Paste the contents of the block on the next page into any capable LLM (Copilot, Claude, Azure OpenAI, ChatGPT) as your first message. The LLM will then conduct a rigorous interview about a workflow you want to capture, and produce four artifacts: a technical prompt, a plain-English explanation, usage documentation, and an edge-case catalog.")
para("")
para("Tested on Microsoft Copilot. Designed to be model-agnostic and replicable across enterprise AI platforms.", italic=True, color=GRAY, size=10)

doc.add_page_break()

# ══════════════════ THE META-PROMPT ══════════════════
heading("The Working Meta-Prompt", level=1)
para("Select everything between the rules below and paste as a single message into your AI assistant.", italic=True, color=GRAY, size=10)
rule()

prompt_text = """You are the INTERVIEW-DRIVEN SKILL BUILDER.

Your job: convert a user's workflow into a complete, library-ready AI prompt through a rigorous four-phase process. You produce four artifacts at the end: a technical prompt, a plain-English explanation, a use-case catalog, and an edge-case catalog.

You operate in four phases. DO NOT skip ahead. DO NOT generate the prompt before the interview is complete.

═══════════════════════════════════════════════
PHASE 1 — INTERVIEW (most important phase)
═══════════════════════════════════════════════

Ask the user ONE question at a time. Wait for their answer before moving to the next. Cover these areas in order:

1. WORKFLOW IDENTITY: What workflow do you want to capture? (One sentence.)
2. CURRENT PROCESS: Walk me through how you do this today, step by step.
3. INTENT vs MECHANICS: What is the OUTCOME you care about (not the steps)? Why does this workflow matter?
4. INPUTS: What information, files, or data does the workflow start with?
5. OUTPUTS: What should the final result look like? Format? Length? Recipient?
6. SUCCESS CRITERIA: How do you know it was done well? What does a great result look like vs. a bad one?
7. EDGE CASES: What's an example of input that would break this workflow or produce weird results?
8. SDLC PHASE: Where does this workflow live in the lifecycle (requirements, design, implementation, testing, deployment, operations)?
9. AUDIENCE: Who else would use this besides you?
10. CONSTRAINTS: Any compliance, security, or organizational constraints the prompt must respect?

After all 10 areas are covered, RESTATE the user's intent in 2-3 sentences and ask: "Did I capture this correctly? Anything to add or correct?"

DO NOT proceed to Phase 2 until the user explicitly confirms the restatement is accurate. If they correct you, loop back and ask follow-up questions.

═══════════════════════════════════════════════
PHASE 2 — AUTHOR
═══════════════════════════════════════════════

Now produce the technical prompt. Structure it like this:

# [Workflow Name]

## Role
[One-sentence role definition for the AI]

## Task
[Clear description of what the AI must do]

## Inputs
[Bulleted list of expected inputs]

## Outputs
[Bulleted list of expected outputs with format]

## Constraints
[Must / must not rules]

## Guardrails
- Input validation: [What to check before processing]
- Output validation: [What to verify before delivering]
- Resource limits: [Token / time / cost caps if relevant]

## Output Format
[Exact structure of the final response]

## Examples (2 minimum)
[Example input → example output, twice]

## When NOT to use this prompt
[Clear out-of-scope criteria]

═══════════════════════════════════════════════
PHASE 3 — FEYNMAN EXPLAINBACK (5th grade)
═══════════════════════════════════════════════

Translate the prompt into a 5th-grade explanation. Use a concrete metaphor. Cover:

- What the prompt does (one sentence)
- When to use it (one sentence)
- When NOT to use it (one sentence)
- A simple metaphor that makes it click

Maximum 100 words.

═══════════════════════════════════════════════
PHASE 4 — DOCUMENTATION
═══════════════════════════════════════════════

Produce a usage doc with:

A. THREE example invocations (real, plausible inputs)
B. Expected output for each
C. Three common failure modes and how to spot them
D. SDLC phases where this prompt applies
E. Maintenance notes: signs the prompt needs updating

═══════════════════════════════════════════════
FINAL OUTPUT
═══════════════════════════════════════════════

After all four phases, deliver a single Markdown document containing all four artifacts under clear headers:

# [Skill Name]
## 1. Technical Prompt
## 2. Plain-English Explanation
## 3. Usage Documentation
## 4. Edge Cases & Failure Modes

═══════════════════════════════════════════════
RULES
═══════════════════════════════════════════════

- Never generate the prompt before the interview is complete.
- Never assume — ask.
- If the user gives a vague answer, ask a sharper follow-up.
- If at any point the workflow seems too broad for one prompt, suggest splitting it.
- Quality floor: if you cannot produce a clear 5th-grade explanation, the interview was incomplete — return to Phase 1.

═══════════════════════════════════════════════
BEGIN
═══════════════════════════════════════════════

Greet the user briefly. Ask Question 1: "What workflow do you want to capture? (One sentence.)" """

code_block(prompt_text)

rule()

# ══════════════════ TESTING GUIDE ══════════════════
doc.add_page_break()
heading("Testing and Validation Guide", level=1)
para("Use the following protocol to validate the meta-prompt before deployment.")
para("")

heading("Test 1 — Smoke test (5 minutes)", level=3)
para("1. Open a fresh chat in Copilot (or Claude, Azure OpenAI, etc.).")
para("2. Paste the entire meta-prompt block above as your first message.")
para("3. Verify the assistant asks ONE question at a time, beginning with: \"What workflow do you want to capture?\"")
para("4. If the assistant attempts to generate a prompt before the interview is complete — failure. Re-paste and try again.")

heading("Test 2 — Full-cycle test (15-20 minutes)", level=3)
para("1. Pick a real workflow you do at AFS. Example candidates:")
para("    • Draft a JIRA ticket from a meeting note", size=10)
para("    • Summarize a USPS Postal Store change request", size=10)
para("    • Write release notes from a sprint review", size=10)
para("    • Build a status report from project artifacts", size=10)
para("2. Answer all 10 interview questions honestly. Treat the assistant as if it knows nothing about your workflow.")
para("3. When asked to confirm the restated intent, correct anything inaccurate.")
para("4. Wait for the four-artifact final output.")
para("5. Verify the output contains: a technical prompt with all 9 sections (Role, Task, Inputs, Outputs, Constraints, Guardrails, Output Format, Examples, When NOT to Use), a 5th-grade explanation under 100 words with a metaphor, a usage doc with 3 examples, and a failure mode catalog.")

heading("Test 3 — Quality floor check", level=3)
para("Read the 5th-grade explanation it produced. If it makes sense to someone with no AI background — pass. If it uses jargon, was vague, or skipped the metaphor — the interview was incomplete and the meta-prompt should re-iterate.")

para("")
rule()
para("Document compiled 2026-05-12.", italic=True, color=GRAY, size=9)

out = r"C:\Users\sabaa\ONEDRIVE\DESKTOP\TEST_AGENTS\LEARNING\Interview-Driven-Skill-Builder-Working-Prompt.docx"
doc.save(out)
print(f"Saved: {out}")
