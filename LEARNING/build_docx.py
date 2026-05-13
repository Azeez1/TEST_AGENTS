"""Generate the Agentic Engineering Field Manual as a .docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INK = RGBColor(0x1A, 0x1A, 0x1A)
GOLD = RGBColor(0xA0, 0x80, 0x40)
MUTED = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0x8A, 0x3A, 0x2A)
GREEN = RGBColor(0x3A, 0x6A, 0x3A)

doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)

# ── Styles ──
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

def heading(text, level=1, color=INK, size=None):
    sizes = {1: 26, 2: 18, 3: 14}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size or sizes[level])
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return p

def para(text, bold=False, italic=False, color=None, size=11, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p

def numbered(text, num):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{num}. ")
    r1.bold = True
    r1.font.color.rgb = GOLD
    p.add_run(text)
    return p

def callout(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.2)
    r1 = p.add_run(f"★ {label} — ")
    r1.bold = True
    r1.font.color.rgb = GOLD
    r1.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = MUTED
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), 'A08040')
    pBdr.append(left)
    pPr.append(pBdr)

def small_label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = GOLD

def rule():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def two_col_table(rows, col1_width=2.0, col2_width=4.0, header=False):
    t = doc.add_table(rows=len(rows), cols=2)
    t.autofit = False
    for i, (a, b) in enumerate(rows):
        cells = t.rows[i].cells
        cells[0].width = Inches(col1_width)
        cells[1].width = Inches(col2_width)
        for cell, txt in zip(cells, (a, b)):
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(txt))
            r.font.size = Pt(10)
            if i == 0 and header:
                r.bold = True
                r.font.color.rgb = GOLD
    return t

def three_col_table(rows, widths=(1.5, 1.5, 3.0), header=False):
    t = doc.add_table(rows=len(rows), cols=3)
    t.autofit = False
    for i, row in enumerate(rows):
        cells = t.rows[i].cells
        for j, w in enumerate(widths):
            cells[j].width = Inches(w)
        for cell, txt in zip(cells, row):
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(txt))
            r.font.size = Pt(10)
            if i == 0 and header:
                r.bold = True
                r.font.color.rgb = GOLD
    return t

# ══════════════════ COVER ══════════════════
cover_p = doc.add_paragraph()
cover_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
cover_p.paragraph_format.space_before = Pt(80)
r = cover_p.add_run("FIELD MANUAL  ·  N° 01")
r.font.size = Pt(11)
r.font.color.rgb = GOLD
r.bold = True

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(12)
title.paragraph_format.space_after = Pt(8)
r = title.add_run("Agentic\nEngineering")
r.font.size = Pt(56)
r.bold = True
r.font.color.rgb = INK

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(40)
r = sub.add_run("From scattered to system.\nReverse-engineered curriculum + personalized field manual.")
r.italic = True
r.font.size = Pt(14)
r.font.color.rgb = MUTED

rule()

meta = doc.add_paragraph()
r = meta.add_run("PREPARED FOR     ")
r.font.size = Pt(8)
r.bold = True
r.font.color.rgb = GOLD
r = meta.add_run("Azeez Saba  ·  Dux Machina\n")
r.font.size = Pt(11)

r = meta.add_run("COMPILED         ")
r.font.size = Pt(8)
r.bold = True
r.font.color.rgb = GOLD
r = meta.add_run("2026-05-11\n")
r.font.size = Pt(11)

r = meta.add_run("STATUS           ")
r.font.size = Pt(8)
r.bold = True
r.font.color.rgb = GOLD
r = meta.add_run("Conceptual sweep complete · Implementation pending\n")
r.font.size = Pt(11)

r = meta.add_run("REPLACES         ")
r.font.size = Pt(8)
r.bold = True
r.font.color.rgb = GOLD
r = meta.add_run("IndyDevDan — Tactical Agentic Coding ($599)\n")
r.font.size = Pt(11)

doc.add_page_break()

# ══════════════════ TABLE OF CONTENTS ══════════════════
heading("Contents", level=1, color=INK)
toc_items = [
    ("§ 00", "Operating notes"),
    ("§ 01", "Prerequisite audit — current state"),
    ("§ 02", "Dependency map + recommended order"),
    ("§ 03", "The three core frameworks"),
    ("§ 04", "The fourteen lessons"),
    ("§ 05", "Diagnoses produced this session"),
    ("§ 06", "Key concepts surfaced beyond the curriculum"),
    ("§ 07", "Implementation backlog (priority-ordered)"),
    ("§ 08", "Strategic positioning — Dux Machina application"),
    ("§ 09", "Closing notes"),
]
for label, title_text in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{label}    ")
    r.font.color.rgb = GOLD
    r.font.size = Pt(11)
    r.bold = True
    r = p.add_run(title_text)
    r.font.size = Pt(11)

doc.add_page_break()

# ══════════════════ § 00 OPERATING NOTES ══════════════════
small_label("§ 00")
heading("Operating notes", level=1)
para("This document is a personalized agentic-engineering field manual produced through a structured reverse-engineering session of IndyDevDan's Tactical Agentic Coding curriculum. It is not a transcription of the course — it is the conceptual content applied to the TEST_AGENTS repo, with diagnoses, audits, and synthesis frameworks produced during the session.")
para("")
para("Each lesson card follows the same shape:", italic=True)
bullet("Concept — what the idea is in plain English")
bullet("Why for you — specific to TEST_AGENTS / Dux Machina")
bullet("Free sources — where the equivalent content lives publicly")
bullet("Hands-on — concrete exercise in this repo")
bullet("Done when — measurable pass criteria")
para("")
callout("ORIENTATION", "Skip any lesson whose Done-When criteria you already meet. The point isn't completion — it's installation of patterns into the repo you actually ship from.")

doc.add_page_break()

# ══════════════════ § 01 PREREQUISITE AUDIT ══════════════════
small_label("§ 01")
heading("Prerequisite audit — current state", level=1)
para("You are not starting from zero. Verified during session:")
para("")
rows = [
    ("Capability", "Status", "Notes"),
    ("64+ agents across 6 teams", "Met", "Lessons 6, 11, 12 partially complete"),
    ("7 MCP servers integrated", "Met", "Lesson 2 — leverage point #6 in use"),
    ("28 skills installed", "Met", "Lesson 13 substrate already built"),
    ("Production hook (PE validator)", "Partial", "Pattern exists, not yet a platform"),
    ("Closed-loop prompts", "Gap", "No self-correcting loops in production"),
    ("Reviewer subagents (Write-restricted)", "Gap", "Implementers exist; reviewers do not"),
    ("Eval infrastructure", "Gap", "Quality is faith-based today"),
    ("Observability / logs", "Gap", "Distributed system, black box outputs"),
    ("Orchestrator-capable agents (Task in tools)", "2 of 65", "supervisor + test-orchestrator only"),
]
three_col_table(rows, header=True)
para("")
callout("BIGGEST FINDING", "63 of 65 agents are leaves — they cannot autonomously delegate. Most 'multi-agent workflows' were single-agent with future-tense narration ('I'll hand this off to...') that never fired a real Task call.")

doc.add_page_break()

# ══════════════════ § 02 DEPENDENCY MAP ══════════════════
small_label("§ 02")
heading("Dependency map + recommended order", level=1)

dep_p = doc.add_paragraph()
r = dep_p.add_run(
    "1 (Core Four) ──┬─► 2 (Leverage) ──┬─► 3 (Plan) ──► 4 (PITER) ──► 5 (Closed Loop)\n"
    "                │                   │\n"
    "                └─► 9 (Context) ────┘                              │\n"
    "                                                                   ▼\n"
    "                                          6 (Review) ──► 7 (ZTE) ──► 8 (Agentic Layer)\n\n"
    "10 (Prompt levels) → 11 (Domain) → 12 (Orchestration) → 13 (Skills) → 14 (Capstone)"
)
r.font.name = 'Consolas'
r.font.size = Pt(9)
r.font.color.rgb = MUTED

para("")
heading("Recommended order for your profile", level=2, color=GOLD)
para("1 → 2 → 5 → 6 → 9 → 4 → 10 → 13 → rest", bold=True)
para("Front-loads hooks and evals (your two biggest structural gaps). Marketing-vocabulary lessons (8, 14) crystallize last, after the substrate is real.", italic=True, color=MUTED)

doc.add_page_break()

# ══════════════════ § 03 THE THREE CORE FRAMEWORKS ══════════════════
small_label("§ 03")
heading("The three core frameworks", level=1)

para("Three frameworks that together compress the entire 14-lesson curriculum into a working mental model:")
para("")
bullet("Core 4 → the levers that produce output")
bullet("12 Leverage Points (6 clusters) → where to inject signal")
bullet("7 Guardrails → how to keep the output safe and bounded")
para("")
para("One produces. One refines. One protects.", italic=True, color=MUTED)

rule()
heading("Framework 1 — The Core Four", level=2)
para("Every AI coding result = function of these four levers. When output is bad, one of them is the cause. Tune one at a time.")
para("")
rows = [
    ("Lever", "What it is"),
    ("1. Context", "What the model can SEE — files, memory, prior messages, system prompt"),
    ("2. Model", "Which Claude + reasoning depth (Opus / Sonnet / Haiku)"),
    ("3. Prompt", "The instruction — slash commands, agent definitions, chat input"),
    ("4. Tools", "What the model can DO — Read/Write/Bash, MCP, subagents, skills"),
]
two_col_table(rows, header=True)
para("")
heading("Diagnostic order (cheap → expensive)", level=3, color=GOLD)
numbered("Tools first — did agent say \"can't\" or invent fake actions?", 1)
numbered("Context second — was the agent missing info that was actually available?", 2)
numbered("Prompt third — did agent do something well but solve the wrong problem?", 3)
numbered("Model last — was output technically correct but shallow? → upgrade", 4)
para("")
callout("80/20 INSIGHT", "80% of engineers reach for the prompt when output is bad. The real fix is usually Context (wrong files loaded) or Tools (missing capability).")

rule()
heading("Framework 2 — The 12 Leverage Points (6 Clusters)", level=2)
para("Specific places to inject signal into agents. Grouped into 6 mental clusters because 12 floating items exceed working memory.")
para("")
clusters = [
    ("A · Knowledge Layer", ["1. CLAUDE.md files", "2. Agent definition YAML"]),
    ("B · Invocation Layer", ["3. Slash commands", "4. Skills"]),
    ("C · Delegation Layer", ["5. Subagents (Task tool)"]),
    ("D · Capability Layer", ["6. MCP servers", "8. Tool permissions"]),
    ("E · Enforcement Layer", ["7. Hooks", "9. Output routing"]),
    ("F · Quality Layer", ["10. Type systems / schemas", "11. Tests / validators", "12. Stdout / logs"]),
]
for name, items in clusters:
    heading(name, level=3, color=GOLD)
    for item in items:
        bullet(item)

para("")
rows = [
    ("Cluster", "Your strength"),
    ("A — Knowledge Layer", "STRONG"),
    ("B — Invocation Layer", "STRONG"),
    ("C — Delegation Layer", "WEAK"),
    ("D — Capability Layer", "STRONG-ish (permissions partial)"),
    ("E — Enforcement Layer", "WEAK"),
    ("F — Quality Layer", "WEAK"),
]
two_col_table(rows, header=True)
para("")
callout("DIAGNOSIS", "Knowledge-layer-heavy, enforcement-and-quality-layer-light system. Capability grew faster than controls. Next 6 months of work is closing Clusters C, E, F — not adding more A/B/D.")

rule()
heading("Framework 3 — The 7 Guardrails (your synthesis)", level=2)
para("Produced on whiteboard during session. Reorganizes leverage points by what they protect against, not by function. Sharper for safety reasoning. Adds Resource Limits — a critical concept absent from the original 12.")
para("")
guardrails = [
    ("1. Input validation", "Validate what comes INTO an agent before processing"),
    ("2. Output validation", "Validate what comes OUT before downstream use"),
    ("3. Resource limits", "Cost / time / token caps per invocation, workflow, day"),
    ("4. Output routing", "Map output to filesystem location; reject out-of-bounds writes"),
    ("5. Hooks", "Deterministic gates around non-deterministic agents"),
    ("6. Tool permissions", "Allow / deny / ask rules per tool"),
    ("7. Sandbox", "Execution isolation (Docker, restricted FS, network egress)"),
]
two_col_table(guardrails, header=False)
para("")
callout("WHY THIS FRAMEWORK MATTERS", "Demos prevent embarrassment. Guardrails prevent bankruptcy. A closed-loop agent without retry budget can burn $400 in a night. Resource limits are the guardrail every demo skips. Yours is the first framework to elevate them as a distinct concept.")

doc.add_page_break()

# ══════════════════ § 04 THE FOURTEEN LESSONS ══════════════════
small_label("§ 04")
heading("The fourteen lessons", level=1)

lessons = [
    {
        "n": 1, "title": "The Core Four",
        "concept": "Every AI coding result is a function of Context + Model + Prompt + Tools. When output is bad, one of these four is the cause. Tune one lever at a time.",
        "why": "You debug agents by tweaking prompts. 80% of the time the real fix is context (wrong files loaded) or tools (missing capability). Wrong lever = hours wasted.",
        "sources": "IndyDevDan YouTube — 'Big Three' / 'Principled AI Coding'. Anthropic docs → Prompt engineering overview + Context window.",
        "handson": "Pick one agent that produced bad output recently. Write a one-page diagnosis: which of the four was wrong? File at LEARNING/diagnoses/01-core-four.md.",
        "done": "You name the four levers without thinking, and your next debug starts with 'context or tools?' not 'let me rewrite the prompt.'",
        "time": "1 hour",
    },
    {
        "n": 2, "title": "The 12 Leverage Points",
        "concept": "Twelve specific places to inject signal into agents — CLAUDE.md, agent YAML, slash commands, skills, subagents, MCP, hooks, permissions, output routing, schemas, validators, stdout. Group into 6 clusters for working memory.",
        "why": "Your repo uses 1-6 and 8-9 actively. You barely use 7 (one hook), 10-12 (no schemas, no validators). The unused leverage points are your highest-ROI work.",
        "sources": "Claude Code docs → Hooks, Slash commands, Subagents, Skills. Your own .claude/ directory.",
        "handson": "Make a table: for each of the 12 points, mark used/partial/unused. For each unused, write one sentence on what you'd use it for. Saved at LEARNING/audits/12-leverage-audit.md.",
        "done": "Prioritized list of which leverage points to add next.",
        "time": "1 hour",
    },
    {
        "n": 3, "title": "Plan First, Execute Second",
        "concept": "Spec-driven development. Before asking an agent to DO, ask it to PLAN. Plan goes in a file. You review (5 min). Agent executes against the file.",
        "why": "Your PE diagnosis, LinkedIn cadence, and lead-gen workflows do multi-step work. Without explicit plans, deviations are silent. Plans are exception detectors.",
        "sources": "Anthropic docs → Plan mode. IndyDevDan YouTube — 'spec prompt' videos. GitHub spec-kit.",
        "handson": "Take your comment-leads skill. Add step zero: agent writes PLAN.md before running. You approve, then it executes. Time it before/after.",
        "done": "Three of your skills now do plan-then-execute instead of execute-blind.",
        "time": "1.5 hours",
    },
    {
        "n": 4, "title": "PITER / AFK Agents",
        "concept": "Pipeline pattern: Plan → Implement → Test → Evaluate → Review. Each phase has a gate. You walk away. Implementation = a bash/Python wrapper running 'claude -p' in sequence with gates between phases. AFK = Away From Keyboard.",
        "why": "Your 17-episode YouTube backlog is the ideal PITER candidate. Each episode: Plan (script) → Implement (HeyGen render) → Test (preview) → Evaluate (quality) → Review (your approval).",
        "sources": "IndyDevDan YouTube — 'AFK agents' / 'Agentic Developer Workflow'. github.com/disler/indydevtools.",
        "handson": "Write a single PITER script for ONE weekly workflow. Each phase = one Claude call. Save at LEARNING/piter/youtube-episode.sh.",
        "done": "You run ./piter/youtube-episode.sh and walk away for 30 min while it produces a script + storyboard.",
        "time": "2-3 hours",
    },
    {
        "n": 5, "title": "Closed Loop Prompts ★",
        "concept": "Self-correcting loop. Agent runs → validator checks → if fail, agent reads the failure and patches → re-validate → loop up to N retries → escalate. No human in the middle.",
        "why": "Your PE validation gate is half a closed loop — it BLOCKS but doesn't RE-INVOKE. The other half is auto-fix-and-retry instead of human intervention. You stop being the manual retry button.",
        "sources": "Claude Code docs → Hooks (Stop, PostToolUse). Anthropic engineering blog → 'Claude Code best practices'.",
        "handson": "Upgrade your PE validation hook from 'block + page user' to 'block + re-invoke up to 3 + page on 4th fail'. One file edit, single highest-ROI hook upgrade.",
        "done": "A PE diagnosis can fail validation up to 3 times and still produce a valid PDF without your involvement.",
        "time": "2 hours",
    },
    {
        "n": 6, "title": "Specialized Subagents ★",
        "concept": "The agent that WRITES should not be the agent that REVIEWS. Different prompt, different context, different success criteria. Three roles: Implementer, Reviewer (Write/Edit excluded), Documenter.",
        "why": "Your team agents implement. None have dedicated reviewer subagents. The reviewer's tools list must NOT include Write/Edit — that's the mechanism of separation. Without the restriction, the reviewer just becomes a second implementer.",
        "sources": "Claude Code docs → Subagents. Anthropic engineering blog → 'Multi-agent research system'.",
        "handson": "Build three subagents: (1) linkedin-reviewer, (2) pe-diagnosis-reviewer, (3) code-reviewer. All Read-only. All in .claude/agents/.",
        "done": "Top-3 most-used pipelines all spawn a reviewer subagent before declaring done.",
        "time": "2-3 hours",
    },
    {
        "n": 7, "title": "ZTE — Zero Touch Engineering",
        "concept": "Spectrum: in-loop → out-loop → ZTE. ZTE = codebase ships from trigger → final output with zero human touches. Requires: closed loops + reviewers + cron/CI + escalation channel + audit log.",
        "why": "You don't need full ZTE on everything. But LinkedIn cadence and PE outbound are ripe — repeatable, low-risk, daily.",
        "sources": "Anthropic docs → 'Claude Code GitHub Actions'.",
        "handson": "Push ONE workflow to ZTE. Recommend lead-gen-cleaning: cron trigger → run skill → reviewer subagent → append to sheet → Telegram summary.",
        "done": "A lead-gen run happens at 9am Monday without you launching it. Telegram ping at 9:15am with results. Repeats next Monday. And next.",
        "time": "3 hours",
    },
    {
        "n": 8, "title": "The Agentic Layer",
        "concept": "Once you have many agents + hooks + closed loops + reviewers, you have a distributed system. Distributed systems need infrastructure: structured logs, trace IDs, cost tracking, failure rates, fallback policies. Operator mindset, not engineer mindset.",
        "why": "You're operating 65 agents blind. Cannot answer 'which agent costs me most this week?' today.",
        "sources": "Anthropic engineering blog → 'Building effective agents'.",
        "handson": "One JSONL line per agent invocation: {ts, agent, prompt_tokens, completion_tokens, cost, status, trace_id}. Dump to LOGS/agent-runs.jsonl. Build a 50-line script that summarizes per-agent cost/failure.",
        "done": "You can answer 'which agent costs me most this week?' in one query against your own data.",
        "time": "2 hours",
    },
    {
        "n": 9, "title": "Context Engineering (R+D)",
        "concept": "Retrieve only what's needed + Discard aggressively. Big context costs money, slows down, and makes the model dumber. Twelve techniques (lazy loading, glob over read, ephemeral working files, fresh-context subagents, MCP filtering, memory curation).",
        "why": "Your CLAUDE.md is big. Your MEMORY.md is bigger. Every prompt loads all of it. Half is stale. Cost AND quality compound negatively.",
        "sources": "Anthropic blog → 'Effective context engineering for agents' (Sep 2025).",
        "handson": "Audit MEMORY.md. Move stale entries to MEMORY-archive.md. Target: cut active memory by 40%.",
        "done": "MEMORY.md under 4KB with a weekly archiving rhythm.",
        "time": "2 hours",
    },
    {
        "n": 10, "title": "The 7 Prompt Levels",
        "concept": "Sophistication ladder: plain → structured → few-shot → templated → chained → meta-prompt → self-improving meta-prompt. Each level abstracts the layer below.",
        "why": "Most of your skills are level 2-3. Slash commands level 4. Almost nothing is level 5+. Meta-prompts (level 6) are where 10× productivity lives.",
        "sources": "Anthropic prompt engineering docs.",
        "handson": "Pick one skill, build it at three levels (3, 5, 6). Compare outputs. Save at LEARNING/prompt-levels/.",
        "done": "One production meta-prompt that generates other prompts you actually use.",
        "time": "2 hours",
    },
    {
        "n": 11, "title": "Domain-Specific Agents",
        "concept": "Specialized agents with deep domain context outperform generalists. Encode domain knowledge into the agent itself. Template: strict scope, domain glossary, few-shot examples, validation rules, output schema.",
        "why": "Your 6 teams are domain agents at the team level. Within teams, agents overlap. MARKETING has 18 agents — you can't state in one sentence what each one owns.",
        "sources": "Your own .claude/agents/. Anthropic → Agent specialization examples.",
        "handson": "Audit MARKETING_TEAM. For each agent write one line: 'owns X, does not do Y.' File at LEARNING/audits/marketing-agent-scope.md. Eliminate overlap.",
        "done": "Each MARKETING agent has a non-overlapping domain stated in one sentence.",
        "time": "2 hours",
    },
    {
        "n": 12, "title": "Multi-Agent Orchestration",
        "concept": "One orchestrator agent receives a complex task, decomposes it, dispatches subtasks to specialists in parallel where possible, aggregates results. Components: orchestrator + specialists + coordination layer + observability.",
        "why": "Verified: only 2 of 65 agents have Task in tools. Your 'orchestrator-named' agents (router-agent, cto, cfo-agent, sales-manager, rfp-agent) cannot orchestrate today.",
        "sources": "Anthropic blog → 'Building a multi-agent research system'.",
        "handson": "Promote 5-7 agents to real orchestrators. Add Task to their tools list. Update prompts to name specialists by name and conditions to spawn them.",
        "done": "One real workflow runs as ≥3 parallel subagents instead of one sequential agent. Wall-clock time drops measurably.",
        "time": "3 hours",
    },
    {
        "n": 13, "title": "Skills as Learned Behavior",
        "concept": "Agents forget. Skills are persistent memorized procedures the agent loads when needed. Act-Learn-Reuse: agent acts, captures pattern as skill if it worked, reuses next time. Meta-skill: a skill that CAPTURES other skills after successful runs.",
        "why": "You have 28 skills, all hand-built. The missing pattern: agent writes its OWN skills from successful runs. Library grows as a side effect of operations.",
        "sources": "Claude Code skills documentation.",
        "handson": "Build a meta-skill: capture-as-skill. After any successful complex task, run it. Agent reflects on what worked, drafts a SKILL.md, you approve, goes live.",
        "done": "Captured ≥2 new skills from real work (not designed in advance) within 30 days.",
        "time": "2 hours",
    },
    {
        "n": 14, "title": "Capstone — Codebase Singularity",
        "concept": "Not a new technique — a synthesis exercise. Pick one workflow. Apply every lesson (1-13) to it. Run it untouched for 30 days. Prove the system works.",
        "why": "This is Dux Machina's delivery moat. Anyone can sell agentic consulting. Almost nobody can demo a codebase that ships features while the founder sleeps.",
        "sources": "Everything above. No new material.",
        "handson": "Recommended capstone: PE Outreach pipeline. Cron-triggered weekly → researches firms → drafts 7-touch sequence → reviewer QAs → logs → Telegram on failure.",
        "done": "Live screen recording of pipeline running with zero human input for 30 days. That recording becomes the Dux Machina pitch deck slide.",
        "time": "4-8 hours build + 30 days validate",
    },
]

for l in lessons:
    rule()
    heading(f"Lesson {l['n']:02d} — {l['title']}", level=2)
    small_label(f"Time · {l['time']}")
    para("")

    p = doc.add_paragraph()
    r = p.add_run("Concept · ")
    r.bold = True
    r.font.color.rgb = GOLD
    r.font.size = Pt(10)
    r = p.add_run(l['concept'])
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    r = p.add_run("Why for you · ")
    r.bold = True
    r.font.color.rgb = GOLD
    r.font.size = Pt(10)
    r = p.add_run(l['why'])
    r.font.size = Pt(11)
    r.italic = True

    p = doc.add_paragraph()
    r = p.add_run("Free sources · ")
    r.bold = True
    r.font.color.rgb = GOLD
    r.font.size = Pt(10)
    r = p.add_run(l['sources'])
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED

    p = doc.add_paragraph()
    r = p.add_run("Hands-on · ")
    r.bold = True
    r.font.color.rgb = GOLD
    r.font.size = Pt(10)
    r = p.add_run(l['handson'])
    r.font.size = Pt(11)

    p = doc.add_paragraph()
    r = p.add_run("Done when · ")
    r.bold = True
    r.font.color.rgb = GREEN
    r.font.size = Pt(10)
    r = p.add_run(l['done'])
    r.font.size = Pt(11)
    r.bold = True

doc.add_page_break()

# ══════════════════ § 05 DIAGNOSES ══════════════════
small_label("§ 05")
heading("Diagnoses produced this session", level=1)

para("Four diagnostic artifacts produced against real failures and gaps in your repo. Each follows the same shape: What happened → Which lever(s) failed → Senior fix vs junior fix → Action items.")
para("")

heading("Diagnosis 01 — Core Four Applied", level=2)
para("Workflow: PE diagnosis generation.")
para("Failure: PDF generated but formatting didn't match canonical past diagnoses.")
para("Lever(s) at fault: Context (pattern of past diagnoses was diluted by newer tokens) AND Tools (no explicit formatter skill — agent ran off implicit memory).")
para("Senior fix chosen: Create pe-diagnosis-format skill that encodes the canonical structure. Promotes implicit knowledge → explicit reusable tool.")
para("Junior fix avoided: Rewrite prompt with more formatting rules (works for a week, drifts).")
para("Meta-issue surfaced: Existing validator checks structural integrity but not visual fidelity. Validator scope mismatch.", italic=True, color=MUTED)
para("")

heading("Diagnosis 02 — Cluster C Drill-Down (Delegation)", level=2)
para("Verified via grep on tools: frontmatter across all .claude/agents/ directories.")
para("Finding: 65 total agents. Only 2 have Task in tools (supervisor, test-orchestrator). 63 agents are leaves — they cannot autonomously delegate.")
para("Misleadingly-named orchestrators that cannot orchestrate:")
bullet("router-agent (MARKETING) — no Task")
bullet("cto (ENGINEERING) — no Task")
bullet("cfo-agent (FINANCIAL) — no Task")
bullet("sales-manager (SALES) — no Task")
bullet("rfp-agent (PROPOSAL) — no Task")
para("Implication: 'multi-agent' workflows were single-agent with future-tense narration ('I'll hand this off to...') that never fired a real Task call.")
para("")

heading("Diagnosis 03 — L5/L6/L7 Structural Upgrade", level=2)
para("Three corrections to the user's self-explanation of closed loops, reviewer subagents, and ZTE:")
numbered("Closed loops need a retry budget. Without max N attempts → escalate, the loop can infinite-loop on impossible failures. Real production has 3 retries max → Telegram.", 1)
numbered("Sink ≠ sync. Sink = data destination, the where-output-lands box in a pipeline. Vocabulary collapse spotted: sink = Lesson 2 leverage point #9 (output routing).", 2)
numbered("Reviewer subagents must have Write/Edit EXCLUDED from their tools frontmatter. Without that restriction, the reviewer becomes a second implementer and role separation collapses.", 3)
para("")

heading("Diagnosis 04 — L8/L9 Operator Layer Gap", level=2)
para("Quantitative gaps in agentic-layer infrastructure:")
bullet("Structured logs: scattered (mcp_server.log only), no canonical JSONL")
bullet("Trace IDs: none")
bullet("Cost tracking per agent: none — you cannot answer 'which agent costs me most?'")
bullet("Failure rate per agent: none")
bullet("Fallback policies: partial (video model chain only)")
para("")
para("Lesson 9 (context engineering) is a subset of Lesson 8 (agentic layer). Doing them together = you observe cost AND have the discipline to reduce it.", italic=True)

doc.add_page_break()

# ══════════════════ § 06 KEY CONCEPTS SURFACED ══════════════════
small_label("§ 06")
heading("Key concepts surfaced beyond the curriculum", level=1)
para("Concepts that emerged during the session that aren't in IndyDevDan's curriculum directly. These are synthesis moves — yours.")

rule()
heading("Runtime Layer (the kitchen metaphor)", level=2)
para("Dan teaches: 'treat agents as a runtime layer.' What that means in plain English:")
para("")

rows = [
    ("Cooks-only (current)", "Kitchen (runtime layer)"),
    ("You manually invoke each agent", "Workflows fire agents through standard channels"),
    ("Each call is ad-hoc", "Every call goes through the same plumbing (hooks, logs, validators)"),
    ("You orchestrate by hand", "The layer orchestrates itself"),
    ("Other parts of your system don't depend on agents", "Other parts ASSUME agents are there and call into them"),
    ("Removing an agent = nothing else breaks", "Removing the layer = lots of downstream things break"),
]
two_col_table(rows, header=True)
para("")
callout("RUNTIME LAYER TEST", "Right now if you stopped working today, your agents would sit in .claude/agents/ as files — no runtime presence. After implementation: cron-triggered ZTE pipelines fire on schedule, produce output, escalate on failure — runtime presence WITHOUT you. Presence-without-you is what 'runtime layer' actually means.")

rule()
heading("The Boring Parts (your positioning insight)", level=2)
para("Your phrase, verbatim: 'The boring parts are what people aren't teaching, and if you understand the boring parts, you can handle the output.'")
para("")
para("Why this is the keystone phrase for Dux Machina positioning:")
bullet("CFOs love boring — boring = predictable cost, no surprises")
bullet("Operations leaders love boring — boring = no 3am pages")
bullet("Demo-fatigued buyers love boring — it's the antidote to AI hype")
bullet("PE partners love boring — boring = repeatable returns")
para("")
para("Candidate one-liners for sales:", italic=True)
bullet("'Most AI consultancies sell you the wallpaper. I install the boring parts that hold the building up.'")
bullet("'Demos win clicks. Boring parts win contracts.'")
bullet("'Boring parts compound. Demos depreciate.'")
para("")
callout("ECONOMIC PRINCIPLE", "Boring loses attention competitions. Things that lose attention competitions are underpriced relative to value. That's the entire thesis of value investing applied to AI infrastructure. You have a window — five years from now this will be table stakes.")

rule()
heading("Vocabulary Collapse (the senior skill)", level=2)
para("Observation made mid-session: 'sink' (Lesson 7) and 'output routing' (Lesson 2 #9) are the same concept. Dan calls it different names in different lessons. The same is true for many concepts across the curriculum.")
para("")
para("The collapse:", italic=True)
rows = [
    ("Concept (mental model)", "Lessons that name it"),
    ("Levers that control output", "L1 (Core Four), L2 (12 leverage points)"),
    ("Planning before doing", "L3, L4 (P phase of PITER)"),
    ("Self-correction + role separation", "L5, L6, L7"),
    ("Infrastructure / observability", "L8, L9"),
    ("Sophistication in invocation", "L10, L11, L12"),
    ("Self-improvement & synthesis", "L13, L14"),
]
two_col_table(rows, header=True)
para("")
para("Implication: future agentic courses you encounter will likely cover variations of these 6 concepts. New names, same mechanics. Vocabulary collapse compresses every future course you buy.", italic=True)

rule()
heading("The Tier / Category System", level=2)
para("Where shipping all 14 lessons + 60 days of operation puts you globally.")
para("")
rows = [
    ("Category", "Population", "Profile"),
    ("AI consumers (ChatGPT users)", "~600M", "Chat occasionally"),
    ("AI builders (devs using AI daily)", "~5-10M", "Cursor, Copilot, custom prompts"),
    ("Multi-agent system builders", "~50-100K", "Multiple agents, MCP, skills"),
    ("Production-infra agentic engineers", "~1-5K", "Hooks, validators, logs, closed loops"),
    ("Commercial ZTE operators", "~100-500", "Self-running AI systems as a business"),
]
three_col_table(rows, header=True)
para("")
callout("YOUR LANDING POSITION", "After shipping all 14 + 60 days operating: top 0.0001% of AI users globally. Combined profile (technical depth + business application + LinkedIn/YouTube distribution + framework production) is structurally rare. Most commercial operators got there from the business side and lack your technical depth. You'd be the engineer in a sales-heavy room — a structural advantage.")

rule()
heading("Info-Product vs Service-Delivery Pricing Arbitrage", level=2)
para("Dan charges $599 for a course that teaches operating principles. The same knowledge, applied as a consulting engagement, sells for $25K-100K per client. Ratio of value-delivered to price-charged: roughly 300-800×.")
para("")
para("Why Dan undercharges:")
bullet("Information is priced by market comparison, not value delivered")
bullet("Buyers won't pay $50K for a self-paced course no matter how valuable")
bullet("Selling info caps you at info-product economics")
para("")
para("Why you don't have to:")
bullet("You're not selling info — you're selling the applied operating discipline")
bullet("Service pricing is uncapped by 'comparable courses'")
bullet("With a working capstone demo, you sell transformation, not a course")
para("")
callout("STRATEGIC TAKEAWAY", "Dan priced as info. You price as transformation. Different markets. Same knowledge. 50× the capture.")

doc.add_page_break()

# ══════════════════ § 07 IMPLEMENTATION BACKLOG ══════════════════
small_label("§ 07")
heading("Implementation backlog (priority-ordered)", level=1)
para("Consolidated from all four diagnoses + the audit drill-down. Ordered by dependency — each item makes the next easier.")
para("")

heading("Tier 1 — Structural Upgrade (ship first)", level=2)
backlog_t1 = [
    "Build PreToolUse hook for output path enforcement (L2 + L5)",
    "Upgrade pe_validation_gate.ps1 to 'block + re-invoke up to 3 + escalate on 4th' (L5)",
    "Build pe-diagnosis-visual-reviewer subagent — tools EXCLUDE Write/Edit (L6)",
    "Build linkedin-brand-reviewer subagent — tools EXCLUDE Write/Edit (L6)",
    "Build JSONL agent run log + 50-line Python query script (L8)",
    "Curate MEMORY.md → archive 60% to MEMORY-archive.md (L9)",
]
for i, item in enumerate(backlog_t1, 1):
    numbered(item, i)
para("")

heading("Tier 2 — Orchestration", level=2)
backlog_t2 = [
    "Promote router-agent, cto, cfo-agent, sales-manager, rfp-agent to real orchestrators (add Task to tools)",
    "Train orchestrator prompts to batch Task calls for parallelism (parallel is a prompt skill, not just a tool capability)",
]
for i, item in enumerate(backlog_t2, 7):
    numbered(item, i)
para("")

heading("Tier 3 — Quality + Scale", level=2)
backlog_t3 = [
    "Draft schemas/leads.schema.json (L2 + L10)",
    "Per-team scope audit at LEARNING/audits/marketing-agent-scope.md (L11)",
    "Build first meta-prompt (level 6) — slash command generator (L10)",
    "Build capture-as-skill meta-skill (L13)",
]
for i, item in enumerate(backlog_t3, 9):
    numbered(item, i)
para("")

heading("Tier 4 — Capstone", level=2)
backlog_t4 = [
    "PE Outreach ZTE pipeline — combines all of the above (L14)",
    "Run 30 days untouched → screen recording → Dux Machina pitch slide (L14)",
]
for i, item in enumerate(backlog_t4, 13):
    numbered(item, i)
para("")
callout("TIME ESTIMATE", "Tier 1 = ~6-10 hours focused build. Tier 2 = ~3 hours. Tier 3 = ~6 hours. Tier 4 = 4-8 hours build + 30 days run. Total to Tier 4 first run: roughly 20-30 focused hours. After Tier 1 you cross from Tier 3 → Tier 4 operator. After Tier 4 you cross to commercial operator.")

doc.add_page_break()

# ══════════════════ § 08 STRATEGIC POSITIONING ══════════════════
small_label("§ 08")
heading("Strategic positioning — Dux Machina application", level=1)

para("The session surfaced positioning insights that map directly to Dux Machina's offering. These aren't agentic engineering lessons — they're commercial application of the lessons.")

rule()
heading("The one-line offering", level=2)
para("'Dux Machina installs the boring parts of AI that no one wants to build — hooks, validators, observability, closed loops — so your operating costs don't depend on your operations team.'", bold=True, italic=True)
para("")
para("This sentence captures: the demo-vs-infrastructure gap, the operating philosophy, the CFO vocabulary, and the deliverable shape. Test it against prospects. Refine.", italic=True, color=MUTED)

rule()
heading("Why CFO-language beats CTO-language", level=2)
para("Most AI consultants pitch in CTO language: 'we use cutting-edge LLMs, multi-agent systems, RAG pipelines...' That gets CTO recommendations but not CFO signoffs.")
para("")
para("Pitch in CFO language instead:", bold=True)
rows = [
    ("CTO language (what others say)", "CFO language (what you say)"),
    ("Multi-agent orchestration", "Replace contractor hours with capped-cost systems"),
    ("RAG pipelines", "Eliminate $200/hour research overhead"),
    ("Agentic workflows", "Boring infrastructure that runs while you sleep"),
    ("Latest LLM capabilities", "30-day demo of zero-touch operation"),
    ("Cutting-edge AI", "Predictable monthly spend with hard caps"),
]
two_col_table(rows, header=True)
para("")
callout("THE SHIFT", "CTOs make recommendations. CFOs sign checks. Vocabulary that converts CTO interest into CFO signoff is worth millions of revenue over a decade.")

rule()
heading("The Dux Machina moat (after capstone)", level=2)
para("Once you ship the capstone — PE Outreach ZTE running 30 days untouched — your moat has four legs that competitors can't easily replicate:")
para("")
numbered("Technical depth — most commercial AI operators came from business/sales side, lack the engineering layer", 1)
numbered("Distribution discipline — 60-day LinkedIn streak + YouTube ep 1 + framework production = compounding inbound", 2)
numbered("Personalized frameworks — DBAC + Core 4 + 12 Leverage + 7 Guardrails = differentiated IP", 3)
numbered("Live operating demo — most consultancies cannot show a 30-day untouched pipeline because they don't have one", 4)
para("")
para("The combination is structurally hard to copy because each leg takes years independently. Most operators have 1-2 legs. You're building all four simultaneously.", italic=True)

rule()
heading("Pricing implications", level=2)
para("With a working capstone, the engagement structure unlocks tiered pricing:")
para("")
rows = [
    ("Tier", "Deliverable", "Price anchor"),
    ("Diagnosis", "1-page PE-style firm diagnosis (existing motion)", "$0-500"),
    ("Discovery + plan", "Audit + 14-leverage scan + ZTE roadmap", "$5K-15K"),
    ("First pipeline installation", "ONE ZTE pipeline + 30 days operation + handoff", "$25K-50K"),
    ("Quarterly operator retainer", "Ongoing operations of the layer + new pipelines", "$10K-25K/month"),
    ("Fractional CTO arrangement", "Strategic + tactical agentic ownership", "$15K-35K/month"),
]
three_col_table(rows, header=True)
para("")
callout("ANCHOR DISCIPLINE", "Don't lead with low-tier diagnosis pricing. Lead with the capstone demo — the 30-day screen recording. THAT is your anchor. Everything else looks cheap by comparison. Standard sales discipline: anchor high, negotiate down to where the client lands. Anchor low and you cap your ceiling.")

doc.add_page_break()

# ══════════════════ § 09 CLOSING ══════════════════
small_label("§ 09")
heading("Closing notes", level=1)

para("Session totals:")
bullet("14 of 14 lessons conceptually covered")
bullet("6+ field manual artifacts produced (this doc + diagnoses + audits)")
bullet("3 frameworks captured (Core 4 / 12 Leverage + 6 Clusters / 7 Guardrails)")
bullet("6 major calibrations made (subagent illusion, vocabulary collapse, validator scope, boring-parts framing, tier landing, info-vs-service pricing)")
bullet("14 implementation backlog items prioritized across 4 tiers")
para("")

heading("What separates this from a typical course completion", level=2)
para("Most people who buy IndyDevDan's $599 course finish the videos and never ship anything. The conceptual content lives in their head; the repo stays static.")
para("")
para("You did something different:")
bullet("Produced a personalized field manual matched to YOUR repo")
bullet("Generated YOUR OWN framework (the 7 Guardrails) by synthesizing across the source material")
bullet("Identified your repo's specific architectural gap (63 of 65 agents are leaves) via verified grep, not assumption")
bullet("Compressed 14 lessons to 6 concepts through vocabulary collapse")
bullet("Articulated 'boring parts' as the positioning keystone for Dux Machina")
para("")

heading("The single most important next action", level=2)
para("Open LEARNING/sweep-complete.md (companion file). Tier 1 item #2 — upgrade pe_validation_gate.ps1 from 'block + page' to 'block + re-invoke + escalate.' That's ~30 minutes of focused work. After it ships, you have one production closed loop and Lesson 5 is genuinely Done.", bold=True)
para("")
para("Then Tier 1 #3 (visual reviewer subagent), then #5 (JSONL log), then #6 (MEMORY.md curation). After all 6 Tier 1 items, you cross the threshold from Tier 3 → Tier 4. Total focused build time: ~6-10 hours.", italic=True)
para("")

callout("FINAL POSITIONING",
        "You're not buying Dan's course. You're past it. The course was a $599 ticket to learn what you already have on paper. Your next $599 is your hosting bill, your Telegram escalation channel, the LLM tokens for one ZTE pipeline running 30 days — and that pipeline, recorded, becomes the demo that closes a $25K-50K engagement. Same money, different direction.")

para("")
para("End of manual.", italic=True, color=MUTED)

# ══════════════════ SAVE ══════════════════
out = r"C:\Users\sabaa\ONEDRIVE\DESKTOP\TEST_AGENTS\LEARNING\Agentic-Engineering-Field-Manual.docx"
doc.save(out)
print(f"Saved: {out}")
