"""
Note Parser Module
Handles parsing of raw meeting notes from various formats (PDF, text, docx)
"""

import PyPDF2
import os
from typing import Dict, List, Optional
import pandas as pd
import openpyxl

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class NoteParser:
    """Parse raw meeting notes into structured data"""

    def __init__(self):
        self.supported_formats = ['.pdf', '.txt', '.md', '.docx', '.xlsx', '.xls']

    def parse_file(self, file_path: str) -> str:
        """
        Parse a file and return its content as text

        Args:
            file_path: Path to the notes file

        Returns:
            Raw text content from the file
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext == '.pdf':
            return self._parse_pdf(file_path)
        elif file_ext in ['.txt', '.md']:
            return self._parse_text(file_path)
        elif file_ext == '.docx':
            return self._parse_docx(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return self._parse_excel(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported: {self.supported_formats}")

    def _parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")

        return text.strip()

    def _parse_text(self, file_path: str) -> str:
        """Read plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            raise Exception(f"Error reading text file: {str(e)}")

    def _parse_docx(self, file_path: str) -> str:
        """Extract text from Word document"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx library not installed. Run: pip install python-docx")

        try:
            doc = docx.Document(file_path)
            text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)

            return "\n".join(text).strip()

        except Exception as e:
            raise Exception(f"Error parsing Word document: {str(e)}")

    def _parse_excel(self, file_path: str) -> str:
        """
        Intelligently parse Excel file with structure detection
        Handles: structured data, free-form text, existing user stories, multiple sheets
        """
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            all_content = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)

                # Skip empty sheets
                if df.empty:
                    continue

                all_content.append(f"=== Sheet: {sheet_name} ===\n")

                # Strategy 1: Check if this looks like existing user stories
                if self._is_user_story_format(df):
                    all_content.append("[DETECTED: Existing User Stories - Will reformat/enhance]\n")
                    all_content.append(self._parse_user_story_excel(df))

                # Strategy 2: Check if this is a structured requirements table
                elif self._is_structured_requirements(df):
                    all_content.append("[DETECTED: Structured Requirements Table]\n")
                    all_content.append(self._parse_structured_excel(df))

                # Strategy 3: Free-form text extraction (fallback)
                else:
                    all_content.append("[DETECTED: Free-form content]\n")
                    all_content.append(self._parse_freeform_excel(df))

                all_content.append("\n")

            return "\n".join(all_content).strip()

        except Exception as e:
            raise Exception(f"Error parsing Excel file: {str(e)}")

    def _is_user_story_format(self, df: pd.DataFrame) -> bool:
        """Detect if DataFrame contains existing user stories"""
        columns_lower = [str(col).lower() for col in df.columns]

        # Check for user story indicators
        story_indicators = ['user_story', 'user story', 'story', 'as a', 'i want']
        ac_indicators = ['acceptance_criteria', 'acceptance criteria', 'ac', 'criteria']

        has_story = any(indicator in ' '.join(columns_lower) for indicator in story_indicators)
        has_ac = any(indicator in ' '.join(columns_lower) for indicator in ac_indicators)

        # Also check first few rows for "As a" pattern
        if not has_story and len(df) > 0:
            first_col_text = ' '.join(df.iloc[:, 0].astype(str).head(3).tolist()).lower()
            if 'as a' in first_col_text and 'i want' in first_col_text:
                has_story = True

        return has_story or has_ac

    def _is_structured_requirements(self, df: pd.DataFrame) -> bool:
        """Detect if DataFrame is a structured requirements table"""
        # Must have at least 2 columns and look table-like
        if len(df.columns) < 2:
            return False

        columns_lower = [str(col).lower() for col in df.columns]

        # Common requirement table column indicators
        req_indicators = [
            'feature', 'requirement', 'description', 'functionality',
            'epic', 'module', 'component', 'priority', 'status',
            'user type', 'acceptance', 'notes', 'rationale'
        ]

        matches = sum(1 for col in columns_lower for indicator in req_indicators if indicator in col)

        # If 2+ columns match common requirement terms, it's likely structured
        return matches >= 2

    def _parse_user_story_excel(self, df: pd.DataFrame) -> str:
        """Parse Excel containing existing user stories"""
        output = []

        # Find the user story column
        story_col = None
        ac_col = None

        for col in df.columns:
            col_lower = str(col).lower()
            if any(indicator in col_lower for indicator in ['user_story', 'user story', 'story']):
                story_col = col
            if any(indicator in col_lower for indicator in ['acceptance_criteria', 'acceptance criteria', 'ac', 'criteria']):
                ac_col = col

        # If no explicit column found, assume first column is user stories
        if story_col is None and len(df.columns) > 0:
            story_col = df.columns[0]

        output.append("Existing User Stories Found:\n")

        for idx, row in df.iterrows():
            if pd.notna(row[story_col]) and str(row[story_col]).strip():
                output.append(f"\nStory {idx + 1}: {row[story_col]}")

                if ac_col and pd.notna(row[ac_col]):
                    output.append(f"Current ACs: {row[ac_col]}")

                # Include other relevant columns
                for col in df.columns:
                    if col not in [story_col, ac_col] and pd.notna(row[col]):
                        output.append(f"{col}: {row[col]}")

        return "\n".join(output)

    def _parse_structured_excel(self, df: pd.DataFrame) -> str:
        """Parse structured requirements table"""
        output = []

        # Add column headers
        output.append("Structured Requirements:\n")
        output.append("Columns: " + " | ".join([str(col) for col in df.columns]) + "\n")

        # Parse each row as a requirement
        for idx, row in df.iterrows():
            # Filter out completely empty rows
            row_values = [str(val) for val in row.values if pd.notna(val) and str(val).strip()]
            if not row_values:
                continue

            output.append(f"\nRequirement {idx + 1}:")
            for col in df.columns:
                if pd.notna(row[col]) and str(row[col]).strip():
                    output.append(f"  {col}: {row[col]}")

        return "\n".join(output)

    def _parse_freeform_excel(self, df: pd.DataFrame) -> str:
        """Parse free-form Excel content (fallback method)"""
        output = []

        output.append("Content:\n")

        # Convert all cells to text, preserving some structure
        for idx, row in df.iterrows():
            row_text = []
            for col in df.columns:
                if pd.notna(row[col]):
                    cell_value = str(row[col]).strip()
                    if cell_value and cell_value != 'nan':
                        # If column name is meaningful, include it
                        col_name = str(col)
                        if not col_name.startswith('Unnamed'):
                            row_text.append(f"{col_name}: {cell_value}")
                        else:
                            row_text.append(cell_value)

            if row_text:
                output.append(" | ".join(row_text))

        return "\n".join(output)

    def parse_text(self, text: str) -> str:
        """
        Parse raw text input (for interactive mode)

        Args:
            text: Raw text notes

        Returns:
            Cleaned text
        """
        return text.strip()


def extract_notes(input_source: str) -> str:
    """
    Main entry point for note extraction

    Args:
        input_source: File path or raw text

    Returns:
        Extracted text content
    """
    parser = NoteParser()

    # Check if input is a file path
    if os.path.exists(input_source):
        return parser.parse_file(input_source)
    else:
        # Treat as raw text
        return parser.parse_text(input_source)
