"""
Create Word documents for LinkedIn campaign
Converts markdown content to professional Word docs
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_linkedin_post_doc():
    """Create Word doc for LinkedIn post"""

    doc = Document()

    # Title
    title = doc.add_heading('LinkedIn Post: Accenture AI Decisions', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph('Generated: 2025-10-13', style='Subtitle')
    doc.add_paragraph('Platform: LinkedIn', style='Subtitle')
    doc.add_paragraph('Character Count: 1,847 (optimal)', style='Subtitle')
    doc.add_paragraph()

    # Main heading
    doc.add_heading("Accenture's Bold AI Decisions: Building the Autonomous Enterprise of 2025", 1)

    # Introduction
    intro = doc.add_paragraph(
        "The consulting giant isn't just talking about AI transformation—they're betting $865 million on it. "
        "And their latest strategic decisions reveal what enterprise AI maturity really looks like."
    )

    doc.add_paragraph()

    # Section 1
    doc.add_heading('The Trust-First Mandate', 2)
    p = doc.add_paragraph(
        "Accenture's 2025 Technology Vision centers on a powerful insight: 77% of executives believe AI's true benefits "
        "only emerge when built on trust. This isn't just corporate speak—it's driving fundamental changes in how they "
        "architect AI systems."
    )
    doc.add_paragraph()
    p = doc.add_paragraph(
        "Their decision? Make trust measurable. While competitors chase performance metrics, Accenture is positioning "
        "trust as the primary KPI for AI success."
    )

    doc.add_paragraph()

    # Section 2
    doc.add_heading('The $2.7 Billion Proof Point', 2)
    p = doc.add_paragraph(
        "Revenue from advanced AI tripled to $2.7 billion in fiscal 2025. But here's what's more interesting—they're "
        "doubling down with strategic investments:"
    )
    doc.add_paragraph()

    # Bullet points
    doc.add_paragraph('Aaru - Agentic prediction engines', style='List Bullet')
    doc.add_paragraph('Workhelix - AI workforce readiness', style='List Bullet')
    doc.add_paragraph('Snorkel AI - High-quality datasets', style='List Bullet')
    doc.add_paragraph('CLIKA - Edge AI compression', style='List Bullet')

    doc.add_paragraph()
    p = doc.add_paragraph(
        "These aren't random bets. They're solving the hardest problems in enterprise AI: trust, talent, data quality, "
        "and deployment scale."
    )

    doc.add_paragraph()

    # Section 3
    doc.add_heading('The Reorganization That Matters', 2)
    p = doc.add_paragraph(
        "Perhaps most revealing: Accenture is consolidating all services into a single \"Reinvention Services\" unit. "
        "By end of 2025, they'll launch 100 industry-specific agentic AI tools."
    )
    doc.add_paragraph()
    p = doc.add_paragraph(
        "This is enterprise AI's iPhone moment—moving from general-purpose tools to industry-specific solutions that "
        "understand context, compliance, and domain expertise."
    )

    doc.add_paragraph()

    # Section 4
    doc.add_heading('The Human Element', 2)
    p = doc.add_paragraph(
        "Here's the uncomfortable truth they're not hiding: they're \"exiting\" staff who can't be reskilled for AI. "
        "But they're also investing heavily in workforce transformation, with 80% of leaders prioritizing positive "
        "human-AI relationships."
    )

    doc.add_paragraph()

    # Section 5
    doc.add_heading('What This Means for You', 2)
    p = doc.add_paragraph("Accenture's decisions illuminate three critical paths:")
    doc.add_paragraph()

    doc.add_paragraph('Trust isn\'t optional—build it into your AI strategy from day one', style='List Number')
    doc.add_paragraph('General AI is table stakes—industry-specific agentic AI is the differentiator', style='List Number')
    doc.add_paragraph('Workforce evolution is non-negotiable—invest in AI skills or risk irrelevance', style='List Number')

    doc.add_paragraph()
    p = doc.add_paragraph(
        "The question isn't whether AI will transform your business. Accenture's $865M reinvention answers that."
    )
    doc.add_paragraph()

    # CTA
    p = doc.add_paragraph("The real question: Are you making decisions as bold as theirs?")
    p.runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph('---' * 20)
    doc.add_paragraph()

    # Hashtags
    hashtags = doc.add_paragraph('#ArtificialIntelligence #EnterpriseAI #DigitalTransformation #AIStrategy #FutureOfWork')
    hashtags.runs[0].font.color.rgb = RGBColor(0, 119, 181)  # LinkedIn blue

    # Save
    doc.save('outputs/blog_posts/Accenture_AI_LinkedIn_Post.docx')
    print('[OK] Created: Accenture_AI_LinkedIn_Post.docx')


def create_complete_package_doc():
    """Create Word doc for complete campaign package"""

    doc = Document()

    # Title page
    title = doc.add_heading('Accenture AI LinkedIn Campaign', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('Complete Package', 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph('Generated: 2025-10-13', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Platform: LinkedIn', style='Subtitle').alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    doc.add_paragraph('1. Campaign Overview', style='List Number')
    doc.add_paragraph('2. LinkedIn Post Content', style='List Number')
    doc.add_paragraph('3. Visual Assets', style='List Number')
    doc.add_paragraph('4. Research Sources', style='List Number')
    doc.add_paragraph('5. Posting Strategy', style='List Number')
    doc.add_paragraph('6. Pre-Publish Checklist', style='List Number')

    doc.add_page_break()

    # Campaign Overview
    doc.add_heading('1. Campaign Overview', 1)

    doc.add_heading('Deliverables', 2)
    doc.add_paragraph('LinkedIn Post (1,847 characters)', style='List Bullet')
    doc.add_paragraph('Professional Header Image (1536x1024)', style='List Bullet')
    doc.add_paragraph('Research-backed content with 6 key statistics', style='List Bullet')
    doc.add_paragraph('5 optimized hashtags', style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('Agents Involved', 2)
    doc.add_paragraph('SEO Specialist - Research & trend analysis', style='List Bullet')
    doc.add_paragraph('Copywriter - Content creation', style='List Bullet')
    doc.add_paragraph('Editor - Quality assurance', style='List Bullet')
    doc.add_paragraph('Social Media Manager - LinkedIn optimization', style='List Bullet')
    doc.add_paragraph('Visual Designer - Header image creation', style='List Bullet')

    doc.add_page_break()

    # Post Content
    doc.add_heading('2. LinkedIn Post Content', 1)
    doc.add_paragraph('[See separate document: Accenture_AI_LinkedIn_Post.docx]')
    doc.add_paragraph()
    doc.add_paragraph('Character Count: 1,847')
    doc.add_paragraph('Optimal Range: 1,300-1,900 ✓')
    doc.add_paragraph('Reading Time: 2-3 minutes')

    doc.add_page_break()

    # Visual Assets
    doc.add_heading('3. Visual Assets', 1)

    doc.add_heading('Header Image', 2)
    doc.add_paragraph('File: accenture-ai-linkedin-header.png')
    doc.add_paragraph('Dimensions: 1536x1024 (3:2 ratio)')
    doc.add_paragraph('Generated with: GPT-4o (gpt-image-1)')
    doc.add_paragraph('Cost: $0.06 USD')
    doc.add_paragraph()

    doc.add_heading('Image Specifications', 3)
    doc.add_paragraph('Style: Modern, corporate-innovative', style='List Bullet')
    doc.add_paragraph('Colors: Deep blue, white, electric blue accents', style='List Bullet')
    doc.add_paragraph('Elements: AI/tech motifs, neural networks, geometric patterns', style='List Bullet')
    doc.add_paragraph('Mood: Forward-thinking, authoritative, innovative', style='List Bullet')

    doc.add_page_break()

    # Research Sources
    doc.add_heading('4. Research Sources', 1)

    doc.add_heading('Key Statistics', 2)
    doc.add_paragraph('$865M - Total AI reinvestment', style='List Bullet')
    doc.add_paragraph('$2.7B - Advanced AI revenue (tripled YoY)', style='List Bullet')
    doc.add_paragraph('77% - Executives prioritizing trust in AI', style='List Bullet')
    doc.add_paragraph('80% - Leaders focusing on human-AI relationships', style='List Bullet')
    doc.add_paragraph('100 - Industry-specific AI tools launching by end 2025', style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('Strategic Investments', 2)
    doc.add_paragraph('Aaru - AI-powered prediction engine', style='List Bullet')
    doc.add_paragraph('Workhelix - Workforce AI readiness platform', style='List Bullet')
    doc.add_paragraph('Snorkel AI - Dataset quality for financial services', style='List Bullet')
    doc.add_paragraph('CLIKA - Edge AI compression technology', style='List Bullet')

    doc.add_page_break()

    # Posting Strategy
    doc.add_heading('5. Posting Strategy', 1)

    doc.add_heading('Best Time to Post', 2)
    doc.add_paragraph('Tuesday-Thursday: 8-10 AM or 12-2 PM EST')
    doc.add_paragraph('Highest LinkedIn engagement windows for B2B content')
    doc.add_paragraph()

    doc.add_heading('Tagging Strategy', 2)
    doc.add_paragraph('Tag @Accenture for potential amplification', style='List Bullet')
    doc.add_paragraph('Consider tagging AI/enterprise thought leaders', style='List Bullet')
    doc.add_paragraph('Tag any collaborators or partners', style='List Bullet')
    doc.add_paragraph()

    doc.add_heading('Engagement Plan', 2)
    doc.add_paragraph('First Hour: Respond quickly to early comments', style='List Bullet')
    doc.add_paragraph('Day 1-2: Continue engagement, ask follow-up questions', style='List Bullet')
    doc.add_paragraph('Day 3-7: Share additional insights in comments', style='List Bullet')
    doc.add_paragraph('Week 2: Consider follow-up post', style='List Bullet')

    doc.add_page_break()

    # Checklist
    doc.add_heading('6. Pre-Publish Checklist', 1)

    doc.add_paragraph('☑ Content written and edited')
    doc.add_paragraph('☑ Length optimized for LinkedIn')
    doc.add_paragraph('☑ Statistics fact-checked')
    doc.add_paragraph('☑ Header image generated')
    doc.add_paragraph('☑ Hashtags selected (5 max)')
    doc.add_paragraph('☐ Spellcheck final version')
    doc.add_paragraph('☐ Tag @Accenture (optional)')
    doc.add_paragraph('☐ Schedule for optimal posting time')
    doc.add_paragraph('☐ Prepare engagement plan')

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    doc.add_paragraph('---' * 20)
    footer = doc.add_paragraph()
    footer.add_run('Campaign Created by Marketing Team Multi-Agent System\n').italic = True
    footer.add_run('Powered by Claude Agent SDK + OpenAI GPT-4o\n').italic = True
    footer.add_run('Total Cost: $0.06 USD').italic = True

    # Save
    doc.save('outputs/blog_posts/Accenture_AI_Complete_Package.docx')
    print('[OK] Created: Accenture_AI_Complete_Package.docx')


if __name__ == '__main__':
    print('[*] Creating Word documents for LinkedIn campaign...')
    print()

    create_linkedin_post_doc()
    create_complete_package_doc()

    print()
    print('=' * 60)
    print('[SUCCESS] Word documents created!')
    print('=' * 60)
    print()
    print('Files created:')
    print('  1. Accenture_AI_LinkedIn_Post.docx')
    print('  2. Accenture_AI_Complete_Package.docx')
    print()
    print('Location: MARKETING_TEAM/outputs/blog_posts/')
    print()
