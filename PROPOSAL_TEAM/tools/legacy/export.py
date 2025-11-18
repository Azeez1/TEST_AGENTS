"""
Export modules for generating final deliverables.
"""

import json
from pathlib import Path
from typing import Dict, Optional

from .logger import logger

try:
    from docxtpl import DocxTemplate

    DOCXTPL_AVAILABLE = True
except ImportError:
    DOCXTPL_AVAILABLE = False
    logger.warning("docxtpl not available - DOCX export will be disabled")


class Exporter:
    """Handle export of proposal to various formats."""

    def export_markdown(self, proposal_text: str, output_path: Path):
        """
        Export proposal as Markdown.

        Args:
            proposal_text: Proposal markdown content
            output_path: Output file path
        """
        logger.info(f"Exporting proposal to Markdown: {output_path}")

        output_path.parent.mkdir(parents=True, exist_ok=True)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(proposal_text)

        logger.info(f"Markdown export complete: {len(proposal_text)} characters")

    def export_docx(
        self,
        sections: Dict[str, str],
        output_path: Path,
        template_path: Optional[Path] = None,
    ):
        """
        Export proposal as DOCX using template.

        Args:
            sections: Dictionary of section content
            output_path: Output DOCX file path
            template_path: Optional custom template path

        Note:
            Requires a DOCX template with Jinja2 placeholders.
            If no template provided, creates a basic DOCX.
        """
        if not DOCXTPL_AVAILABLE:
            logger.error("docxtpl not available, cannot export DOCX")
            raise RuntimeError("docxtpl package required for DOCX export")

        logger.info(f"Exporting proposal to DOCX: {output_path}")

        output_path.parent.mkdir(parents=True, exist_ok=True)

        if template_path and template_path.exists():
            # Use provided template
            doc = DocxTemplate(template_path)
            doc.render(sections)
            doc.save(output_path)

        else:
            # Create basic DOCX from markdown
            # This requires python-docx instead of docxtpl
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

                doc = Document()

                # Add title
                title = doc.add_heading(sections.get("rfp_title", "Proposal"), 0)
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Add sections
                for section_name, content in sections.items():
                    if section_name in [
                        "executive_summary",
                        "technical_approach",
                        "management_approach",
                    ]:
                        doc.add_heading(section_name.replace("_", " ").title(), 1)

                        # Split content into paragraphs
                        paragraphs = content.split("\n\n")
                        for para in paragraphs:
                            if para.strip():
                                # Check if it's a heading
                                if para.startswith("#"):
                                    level = min(para.count("#"), 3)
                                    text = para.lstrip("#").strip()
                                    doc.add_heading(text, level)
                                else:
                                    doc.add_paragraph(para.strip())

                doc.save(output_path)

            except ImportError:
                logger.error("python-docx not available for basic DOCX export")
                raise

        logger.info("DOCX export complete")

    def export_json(self, data: Dict, output_path: Path):
        """
        Export data as JSON.

        Args:
            data: Data dictionary
            output_path: Output file path
        """
        logger.info(f"Exporting JSON: {output_path}")

        output_path.parent.mkdir(parents=True, exist_ok=True)

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

        logger.info("JSON export complete")

    def export_bundle(
        self,
        output_dir: Path,
        proposal_md: str,
        requirements_data: Dict,
        compliance_data: Dict,
        qa_report: Dict,
        sections: Optional[Dict[str, str]] = None,
    ):
        """
        Export complete proposal bundle.

        Args:
            output_dir: Output directory
            proposal_md: Proposal markdown
            requirements_data: Requirements JSON
            compliance_data: Compliance matrix data
            qa_report: QA validation report
            sections: Optional section dictionary for DOCX
        """
        logger.info(f"Exporting complete proposal bundle to: {output_dir}")

        output_dir.mkdir(parents=True, exist_ok=True)

        # Export markdown
        self.export_markdown(proposal_md, output_dir / "proposal_draft.md")

        # Export JSON files
        self.export_json(requirements_data, output_dir / "requirements.json")
        self.export_json(compliance_data, output_dir / "compliance_matrix.json")
        self.export_json(qa_report, output_dir / "qa_report.json")

        # Export compliance CSV
        from .compliance import ComplianceMatrixBuilder

        builder = ComplianceMatrixBuilder()
        builder.export_csv(compliance_data, output_dir / "compliance_matrix.csv")

        # Export DOCX if sections provided
        if sections and DOCXTPL_AVAILABLE:
            try:
                self.export_docx(sections, output_dir / "proposal.docx")
            except Exception as e:
                logger.warning(f"DOCX export failed: {e}")

        logger.info("Bundle export complete")

        # Create summary file
        summary = f"""# Proposal Export Summary

**Output Directory:** {output_dir}

**Generated Files:**
- `proposal_draft.md` - Complete proposal in Markdown
- `requirements.json` - Extracted requirements
- `compliance_matrix.json` - Compliance matrix data
- `compliance_matrix.csv` - Compliance matrix spreadsheet
- `qa_report.json` - QA validation report
"""

        if sections and DOCXTPL_AVAILABLE:
            summary += "- `proposal.docx` - Proposal in Word format\n"

        summary += f"""
**Statistics:**
- Total Requirements: {requirements_data.get('metadata', {}).get('total_requirements', 0)}
- MUST/SHALL Requirements: {requirements_data.get('metadata', {}).get('must_count', 0)}
- Compliance Items: {len(compliance_data.get('compliance_items', []))}
- QA Status: {qa_report.get('status', 'UNKNOWN')}
- QA Issues: {len(qa_report.get('issues', []))}
"""

        (output_dir / "SUMMARY.md").write_text(summary)
